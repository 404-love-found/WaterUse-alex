from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/alex-lirio-lucian/develop/WaterUse-alex/LLM_Prompting")
OUT = ROOT / "ODD_EV-parking" / "EV_Parking_Case_Study_Report.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"
BORDER = "A6A6A6"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_row_as_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        set_row_cant_split(row)
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "EV Parking Case Study Brief"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(89, 89, 89)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Page ")
    add_page_number(footer)
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(89, 89, 89)


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Apartment EV Parking Shared Charging Case Study")
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("ODD+D case brief, expected action situations, and current research-design gaps")
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(89, 89, 89)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.add_run(f"Prepared for ODD_EV-parking | {date.today().isoformat()}").font.size = Pt(10)


def add_lead(doc):
    p = doc.add_paragraph()
    run = p.add_run("Current status: ")
    run.bold = True
    p.add_run(
        "the case is conceptually strong enough for pilot AS-extraction tests, but it still needs empirical reference values, "
        "formal queue-policy variants, validation criteria, and a clearer experimental design before it should be treated as a complete ABM study."
    )


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def fill_table(table, rows, widths_dxa, header=True):
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    if header:
        set_row_as_header(table.rows[0])
    for r_idx, row_values in enumerate(rows):
        cells = table.rows[r_idx].cells
        for c_idx, value in enumerate(row_values):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            run.font.size = Pt(9.5)
            if header and r_idx == 0:
                run.bold = True
                set_cell_shading(cells[c_idx], LIGHT_GRAY)
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(value) <= 12 else WD_ALIGN_PARAGRAPH.LEFT


def add_gap_table(doc):
    rows = [
        ["Gap area", "What is missing now", "How to complete it"],
        [
            "Empirical grounding",
            "The ODD lists possible data sources, but no actual reference values are defined for charger count, resident EV ownership, arrival rates, discount size, fault rates, or overstay frequency.",
            "Create a parameter table with baseline, low, and high values. Mark which values come from logs, surveys, literature, or assumptions.",
        ],
        [
            "Queue rule specification",
            "The case says queue fairness is central, but the exact rule is not fully operationalized.",
            "Define FIFO, reservation windows, no-show cancellation, overstay grace period, resident and non-resident eligibility, and how staff override rules.",
        ],
        [
            "Fairness definition",
            "Fair queue access is the core concept, but the normative metric is still broad.",
            "Choose measurable criteria such as wait-time inequality, access share by group, queue-order violations, complaint rate, and completed-session reliability.",
        ],
        [
            "Decision rules",
            "The agent logic is described at a high level, but thresholds and update rules are not yet precise.",
            "Write thresholds for reserving, joining, bypassing, complaining, moving out, and enforcing; define how experience changes probabilities.",
        ],
        [
            "Policy experiments",
            "The case identifies governance mechanisms but does not yet define experiment arms.",
            "Compare baseline FIFO, resident quota, idle-time penalty, stricter enforcement, reservation caps, charger expansion, and maintenance-priority policies.",
        ],
        [
            "Validation plan",
            "There is no current test for whether simulated waiting, utilization, complaints, or violations match plausible real-world behavior.",
            "Use face validation, sensitivity analysis, extreme-condition tests, and comparison with observed session or complaint logs when available.",
        ],
        [
            "Privacy and ethics",
            "The case uses residency status, charger logs, complaints, and possible monitoring, but privacy limits are not specified.",
            "Define what data are anonymized, who can inspect logs, how staff discretion is audited, and how non-resident users are treated fairly.",
        ],
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    fill_table(table, rows, [1900, 3600, 3860])


def add_as_table(doc):
    rows = [
        ["AS", "Expected action situation", "Evaluation meaning"],
        [
            "AS1",
            "Queue-order and prompt move-out compliance between EV users",
            "Captures whether users respect the posted queue and release scarce charging bays promptly.",
        ],
        [
            "AS2",
            "Sequential learning in reservation and off-peak charging behaviour",
            "Captures how visible successful or opportunistic behavior changes later user choices.",
        ],
        [
            "AS3",
            "Shared charger-capacity contribution dilemma among residents",
            "Captures the resident-level dilemma around paying for expanded capacity or improved monitoring.",
        ],
        [
            "AS4",
            "Informal priority exchange between EV user and parking staff",
            "Captures staff discretion and informal access that can weaken queue legitimacy.",
        ],
        [
            "AS5",
            "Formal complaint and enforcement coordination between user and management",
            "Captures whether users report problems and whether management responds with enforcement or repair.",
        ],
        [
            "AS6",
            "Shared charger-occupancy common-pool congestion between EV users",
            "Captures congestion caused by long or unnecessary charger-bay occupation under per-kWh-only pricing.",
        ],
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    fill_table(table, rows, [900, 3900, 4560])


def build_doc():
    doc = Document()
    style_document(doc)
    add_title(doc)
    add_lead(doc)

    doc.add_heading("1. Case Definition", level=1)
    doc.add_paragraph(
        "This original case studies shared electric-vehicle charging in an apartment parking facility. "
        "Residents receive a discounted per-kWh charging rate, non-resident users pay the regular per-kWh rate, "
        "and all users compete for access to a small number of shared charging bays. The research focus is fair queue access rather than tariff optimization."
    )
    doc.add_paragraph(
        "The model setting is a mixed resident and non-resident garage where EV adoption has grown faster than charger capacity. "
        "The resident discount changes demand and perceived entitlement, but it does not formally grant queue priority. "
        "The central governance problem is how scarce charging access is ordered, enforced, and released over time."
    )

    doc.add_heading("2. What the Case Already Has", level=1)
    add_bullets(
        doc,
        [
            "A stable ODD+D description covering purpose, entities, state variables, spatial and temporal scales, learning, sensing, interaction, collectives, heterogeneity, stochasticity, observation, implementation, initialization, input data, and submodels.",
            "A model-facing game scenario that explains mechanisms without listing or numbering the expected AS labels.",
            "A hidden evaluation reference with six expected action situations for manual scoring and future automated evaluation.",
            "A spreadsheet version aligned with the electricity ODD template, with a separate Expected AS sheet so the gold answers are visible but not mixed into the model input.",
        ],
    )

    doc.add_page_break()
    doc.add_heading("3. Main Gaps to Fill", level=1)
    add_gap_table(doc)

    doc.add_heading("4. Expected Action Situations", level=1)
    doc.add_paragraph(
        "These six ASs should be treated as the current gold set for evaluation. They should not be included in prompts used to test whether a model can extract ASs from the ODD and scenario text."
    )
    add_as_table(doc)

    doc.add_heading("5. Recommended Next Steps", level=1)
    steps = [
        "Decide the exact fairness definition before adding more scenario text. The strongest candidates are queue-order compliance, wait-time inequality, and charger-access reliability by user category.",
        "Create a parameter table with baseline values and ranges. This will make the ABM reproducible and easier to defend.",
        "Write policy experiment arms before coding the model. At minimum, compare baseline FIFO, idle-time penalty, stricter enforcement, resident quota, and charger expansion.",
        "Keep the gold AS list outside the model prompt. Use it only for scoring TP, FN, FP, precision, and recall.",
        "After the case text is finalized, run the same 30-run ODD-only and ODD+game extraction protocol used for water and electricity.",
    ]
    for idx, step in enumerate(steps, start=1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(5)
        p.add_run(step)

    doc.add_heading("6. Prompting and Evaluation Boundary", level=1)
    doc.add_paragraph(
        "For model testing, the input should include the ODD text and, when testing ODD+game, the game scenario context. "
        "The Expected AS sheet and scenarios reference should remain hidden from the model. Evaluation should count only true positives, false negatives, false positives, precision, and recall against the six-AS gold set."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
