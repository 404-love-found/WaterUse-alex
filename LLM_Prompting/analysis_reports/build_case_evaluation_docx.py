#!/usr/bin/env python3
"""Build the Word report from generated evaluation tables."""

from __future__ import annotations

from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "analysis_reports" / "model_evaluation_comparison_2026-06-06"
TABLE_DIR = OUT_DIR / "tables"
TODAY = date(2026, 6, 6)


def fmt(value: object, digits: int = 4) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, float):
        return f"{value:.{digits}f}"
    return str(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    tc_pr.append(shade)


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 7.5, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def repeat_table_header(row) -> None:
    """Mark the header row for repetition when Word splits a table across pages."""
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc: Document, df: pd.DataFrame, columns: list[str], widths: list[float], title: str) -> None:
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.allow_autofit = False
    repeat_table_header(table.rows[0])

    for idx, col in enumerate(columns):
        table.columns[idx].width = Inches(widths[idx])
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "F2F4F7")
        set_cell_text(cell, col, bold=True, size=7.5)

    for _, row in df.loc[:, columns].iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(columns):
            set_cell_text(cells[idx], fmt(row[col]), size=7.0)
    doc.add_paragraph("")


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 15, RGBColor(46, 116, 181)),
        ("Heading 2", 12, RGBColor(46, 116, 181)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("ODD AS Extraction Evaluation Tables")
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(31, 36, 48)

    subtitle = doc.add_paragraph(
        f"Date: {TODAY.isoformat()} | Scope: EV-parking, Water-use, Electricity | Metrics: TP, FN, FP, Precision, Recall"
    )
    subtitle.runs[0].font.size = Pt(9)
    subtitle.runs[0].font.color.rgb = RGBColor(111, 118, 138)


def main() -> None:
    doc = Document()
    style_document(doc)
    add_title(doc)

    completeness = pd.read_csv(TABLE_DIR / "run_completeness.csv").fillna("")
    metrics = pd.read_csv(TABLE_DIR / "overall_metrics_by_case_model_condition.csv")
    effects = pd.read_csv(TABLE_DIR / "prompt_effect_by_case_model.csv")
    aggregate = pd.read_csv(TABLE_DIR / "aggregate_metrics_available_cases.csv")

    doc.add_heading("Run Completeness", level=1)
    add_table(
        doc,
        completeness,
        ["Case", "Condition", "Model", "Run_Count", "Expected_Runs", "Missing_Runs", "Duplicate_Run_Count", "Status"],
        [1.0, 0.8, 1.35, 0.65, 0.75, 0.9, 0.9, 0.85],
        "Completeness By Case, Condition, And Model",
    )
    doc.add_page_break()

    doc.add_heading("Overall Evaluation Results", level=1)
    add_table(
        doc,
        metrics,
        ["Case", "Condition", "Model", "Runs", "TP", "FN", "FP", "Precision", "Recall"],
        [0.95, 0.8, 1.25, 0.45, 0.45, 0.45, 0.45, 0.78, 0.78],
        "Evaluation Results By Case, Prompt Condition, And Model",
    )

    doc.add_heading("ODD+game Minus ODD-only", level=1)
    add_table(
        doc,
        effects,
        ["Case", "Model", "Delta_TP", "Delta_FN", "Delta_FP", "Delta_Precision", "Delta_Recall"],
        [0.95, 1.3, 0.65, 0.65, 0.65, 0.95, 0.95],
        "Prompt-Condition Evaluation Deltas",
    )

    doc.add_page_break()
    doc.add_heading("Aggregate Evaluation Results", level=1)
    add_table(
        doc,
        aggregate,
        ["Model", "Condition", "Case_Count", "Runs", "TP", "FN", "FP", "Precision", "Recall"],
        [1.3, 0.8, 0.7, 0.55, 0.5, 0.5, 0.5, 0.78, 0.78],
        "Aggregate Metrics Across Available Case Coverage",
    )

    out_path = OUT_DIR / "ODD_AS_Extraction_Evaluation_Comparison_Report_2026-06-06.docx"
    doc.save(out_path)
    print(out_path)


if __name__ == "__main__":
    main()
